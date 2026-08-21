import sys
try:
    from docx import Document
except ImportError:
    print("python-docx not installed, trying to install...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document

import os

files = [
    ("f:/1/夫/前言.docx", "f:/1/夫/前言.txt"),
    ("f:/1/夫/排盘-【阴盘-阳遁-5局】.docx", "f:/1/夫/排盘-阳遁5局.txt"),
    ("f:/1/夫/排盘-【阴盘-阴遁-5局】.docx", "f:/1/夫/排盘-阴遁5局.txt"),
]

for src, dst in files:
    print(f"Extracting {src} ...")
    doc = Document(src)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                lines.append(" | ".join(row_text))
        lines.append("")
    with open(dst, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Saved to {dst}, lines={len(lines)}")
